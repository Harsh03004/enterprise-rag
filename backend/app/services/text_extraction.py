from io import BytesIO
from pathlib import Path

import fitz
import pytesseract
from docx import Document as DocxDocument
from PIL import Image
from pypdf import PdfReader


PROJECT_ROOT = Path(__file__).resolve().parents[3]


def extract_text(file_path: str) -> str:
    path = PROJECT_ROOT / file_path

    suffix = path.suffix.lower()

    if suffix in {".md", ".txt"}:
        return path.read_text(encoding="utf-8")

    if suffix == ".pdf":
        return extract_pdf(path)

    if suffix == ".docx":
        return extract_docx(path)

    raise ValueError(
        f"Unsupported file type: {suffix}"
    )


def is_usable_text(text: str) -> bool:
    """
    Determine whether extracted text is usable.

    Very short or heavily corrupted text is considered
    unsuitable and will be replaced with OCR output.
    """

    if not text.strip():
        return False

    if len(text.strip()) < 30:
        return False

    alphanumeric_count = sum(
        character.isalnum()
        for character in text
    )

    printable_count = sum(
        character.isprintable()
        for character in text
    )

    if printable_count == 0:
        return False

    alphanumeric_ratio = (
        alphanumeric_count / len(text)
    )

    printable_ratio = (
        printable_count / len(text)
    )

    return (
        alphanumeric_ratio >= 0.45
        and printable_ratio >= 0.85
    )


def extract_pdf(path: Path) -> str:
    """
    Extract text from a PDF.

    Uses normal PDF text extraction when the extracted
    text appears usable. If a page has no text or the
    extracted text appears corrupted/garbled, that page
    is sent through OCR.

    This allows normal, scanned, and mixed PDFs to work.
    """

    reader = PdfReader(path)

    normal_pages = []

    for page in reader.pages:
        text = page.extract_text() or ""
        normal_pages.append(text.strip())

    document = fitz.open(path)

    pages = []

    try:
        for index, normal_text in enumerate(normal_pages):
            if is_usable_text(normal_text):
                pages.append(normal_text)
            else:
                page = document[index]

                ocr_text = ocr_page(page)

                if ocr_text:
                    pages.append(ocr_text)

    finally:
        document.close()

    return "\n\n".join(pages)


def ocr_pdf(path: Path) -> str:
    """
    OCR every page of a PDF.
    """

    document = fitz.open(path)

    pages = []

    try:
        for page in document:
            text = ocr_page(page)

            if text:
                pages.append(text)

    finally:
        document.close()

    return "\n\n".join(pages)


def ocr_page(page: fitz.Page) -> str:
    """
    Render a PDF page to an image and run Tesseract OCR.
    """

    pixmap = page.get_pixmap(
        matrix=fitz.Matrix(2, 2),
        alpha=False,
    )

    image = pixmap.tobytes("png")

    pil_image = Image.open(
        BytesIO(image)
    )

    text = pytesseract.image_to_string(
        pil_image,
        lang="eng",
    )

    return text.strip()


def extract_docx(path: Path) -> str:
    """
    Extract text from a DOCX document.

    Extracts:
    - Normal paragraphs
    - Tables
    - Text contained inside embedded images
    """

    document = DocxDocument(path)

    sections = []

    # ---------------------------------------------------------
    # Paragraphs
    # ---------------------------------------------------------

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            sections.append(text)

    # ---------------------------------------------------------
    # Tables
    # ---------------------------------------------------------

    for table in document.tables:
        rows = []

        for row in table.rows:
            cells = [
                cell.text.strip()
                for cell in row.cells
            ]

            if any(cells):
                rows.append(
                    " | ".join(cells)
                )

        if rows:
            sections.append(
                "\n".join(rows)
            )

    # ---------------------------------------------------------
    # Embedded images
    # ---------------------------------------------------------

    image_text = extract_docx_images(
        document
    )

    if image_text:
        sections.append(
            image_text
        )

    return "\n\n".join(sections)


def extract_docx_images(
    document: DocxDocument,
) -> str:
    """
    Extract embedded images from a DOCX document
    and run OCR on them.
    """

    image_sections = []

    image_index = 0

    for relationship in document.part.rels.values():
        if "image" not in relationship.reltype:
            continue

        image_index += 1

        try:
            image_part = relationship.target_part
            image_bytes = image_part.blob

            image = Image.open(
                BytesIO(image_bytes)
            )

            text = pytesseract.image_to_string(
                image,
                lang="eng",
            ).strip()

            if text:
                image_sections.append(
                    f"[Image {image_index}]\n{text}"
                )

        except Exception as error:
            print(
                f"Failed to OCR DOCX image "
                f"{image_index}: {error}"
            )

    return "\n\n".join(image_sections)